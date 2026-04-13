"""Generate Architecture Document for Tableau to Power BI Migration Tool."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# Brand colors from reference doc
ORANGE = RGBColor(0xD4, 0x6A, 0x12)
DARK_GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GREY = RGBColor(0x77, 0x77, 0x77)
RED = RGBColor(0xCC, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x33, 0x33, 0x33)
TABLE_HEADER_BG = RGBColor(0xD4, 0x6A, 0x12)

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.page_width = Emu(7772400)
section.page_height = Emu(10058400)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)

# --- Style configuration ---
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = BLACK

for level, size in [(1, 16), (2, 14), (3, 12)]:
    style = doc.styles[f'Heading {level}']
    style.font.name = 'Calibri'
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = ORANGE


def add_title(text, size=32, color=ORANGE, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.bold = bold
    r.font.name = 'Calibri'
    return p


def add_subtitle(text, size=18, color=DARK_GREY):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = 'Calibri'
    return p


def add_body(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    # Add grey background via shading
    shading = r._element.get_or_add_rPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F0F0F0'
    })
    shading.append(shd)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
        r.font.name = 'Calibri'
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): 'D46A12'
        })
        shading.append(shd)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10) if cell.paragraphs[0].runs else None

    return table


def add_separator():
    p = doc.add_paragraph()
    r = p.add_run('_' * 70)
    r.font.size = Pt(6)
    r.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)


# ===================================================================
# COVER PAGE
# ===================================================================
doc.add_paragraph()
doc.add_paragraph()
add_title('Architecture Document', size=32, color=ORANGE)
add_subtitle('Tableau to Power BI Migration Tool', size=22, color=DARK_GREY)
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Automated end-to-end migration of Tableau workbook\nsemantic layer to Power BI Desktop')
r.font.size = Pt(12)
r.font.color.rgb = LIGHT_GREY

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Prepared by: USEReady\nDate: April 2026\nVersion: 1.0')
r.font.size = Pt(10)
r.font.color.rgb = LIGHT_GREY

p = doc.add_paragraph()
r = p.add_run('CONFIDENTIAL')
r.font.size = Pt(10)
r.bold = True
r.font.color.rgb = RED

doc.add_page_break()

# ===================================================================
# TABLE OF CONTENTS
# ===================================================================
doc.add_heading('Table of Contents', level=1)
add_separator()

toc_items = [
    '1.  Executive Summary',
    '2.  Architecture Overview',
    '3.  Pipeline Steps - Detailed Walkthrough',
    '4.  Module Reference',
    '5.  DAX Conversion Engine',
    '6.  Data Flow Diagram',
    '7.  Prerequisites & Environment Setup',
    '8.  How to Run the Tool',
    '9.  Output Files & Structure',
    '10. Limitations & Future Enhancements',
    '11. Troubleshooting Guide',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    r = p.runs[0]
    r.font.size = Pt(12)

doc.add_page_break()

# ===================================================================
# 1. EXECUTIVE SUMMARY
# ===================================================================
doc.add_heading('1. Executive Summary', level=1)
add_separator()
doc.add_paragraph()

add_body('The Tableau to Power BI Migration Tool is an automated pipeline that migrates the complete semantic layer from Tableau workbooks (.twbx/.twb) to Power BI Desktop. The tool eliminates the need for manual recreation of data models, calculated fields, measures, relationships, and folder structures.')

add_body('The migration is executed via a single command:')
add_code_block('python migrate.py "path/to/workbook.twbx"')

add_body('The tool performs a 7-step automated pipeline:')
add_bullet('Extracts the Tableau workbook XML from the .twbx package')
add_bullet('Parses all metadata (tables, columns, calculations, parameters, relationships)')
add_bullet('Extracts data from Hyper files and converts to CSV format')
add_bullet('Generates a complete Power BI model (.bim) with M expressions for data loading')
add_bullet('Generates Tabular Editor C# scripts as fallback deployment method')
add_bullet('Creates a PBIP (Power BI Project) and opens it directly in Power BI Desktop')
add_bullet('Outputs a summary of all generated artifacts')

doc.add_paragraph()
add_body('What gets migrated:')

add_table(
    ['Component', 'Tableau', 'Power BI'],
    [
        ['Data Tables', 'Datasources with Hyper extracts', 'Tables with M partitions (CSV)'],
        ['Physical Columns', '<column> and <metadata-record> elements', 'Data columns with correct types'],
        ['Calculated Fields (row-level)', 'IF/CASE/DATEDIFF formulas', 'DAX Calculated Columns'],
        ['Aggregated Calculations', 'SUM/AVG/COUNT formulas', 'DAX Measures'],
        ['LOD Expressions', '{FIXED [Dim]: AGG([Measure])}', 'CALCULATE + ALLEXCEPT'],
        ['Parameters', 'Parameter datasource', 'Parameters table with measures'],
        ['Relationships', 'Datasource blending', 'Many-to-many relationships'],
        ['Display Folders', 'Folders & drill-path hierarchies', 'Column DisplayFolder property'],
    ]
)

doc.add_page_break()

# ===================================================================
# 2. ARCHITECTURE OVERVIEW
# ===================================================================
doc.add_heading('2. Architecture Overview', level=1)
add_separator()
doc.add_paragraph()

add_body('The tool follows a modular architecture with clear separation of concerns:')

doc.add_heading('2.1 High-Level Architecture', level=2)

add_code_block('''
+-------------------+     +------------------+     +-------------------+
|  Input            |     |  Parser Layer    |     |  Generator Layer  |
|  (.twbx/.twb)     | --> |                  | --> |                   |
|                   |     |  xml_parser.py   |     |  bim_generator.py |
|  Hyper files      |     |  extractor.py    |     |  pbi_generator.py |
|                   |     |  dax_converter.py|     |                   |
+-------------------+     +------------------+     +-------------------+
                                                           |
                                                           v
                          +------------------+     +-------------------+
                          |  Output          |     |  Deploy Layer     |
                          |                  | <-- |                   |
                          |  model.bim       |     |  migrate.py       |
                          |  data/*.csv      |     |  (PBIP + TE2 CLI) |
                          |  scripts/*.cs    |     |                   |
                          +------------------+     +-------------------+
''')

doc.add_heading('2.2 Module Responsibilities', level=2)

add_table(
    ['Module', 'File', 'Responsibility'],
    [
        ['Extractor', 'parser/extractor.py', 'Extracts .twb XML from .twbx ZIP archive'],
        ['XML Parser', 'parser/xml_parser.py', 'Parses Tableau XML to extract all metadata (datasources, columns, calculations, joins, relationships, worksheets, parameters, folders)'],
        ['DAX Converter', 'parser/dax_converter.py', 'Converts Tableau formula syntax to DAX (IF/ELSEIF, CASE/WHEN, LOD, DATEDIFF, aggregation functions, field references)'],
        ['Model Builder', 'parser/model_builder.py', 'Assembles all extracted metadata into a unified dictionary'],
        ['BIM Generator', 'parser/bim_generator.py', 'Generates .bim file (TOM JSON) with tables, M partitions, calculated columns, measures, relationships, display folders'],
        ['PBI Generator', 'parser/pbi_generator.py', 'Generates Tabular Editor C# scripts for alternative deployment'],
        ['Orchestrator', 'migrate.py', 'End-to-end pipeline: extract, parse, generate, deploy to PBI Desktop'],
    ]
)

doc.add_page_break()

# ===================================================================
# 3. PIPELINE STEPS
# ===================================================================
doc.add_heading('3. Pipeline Steps - Detailed Walkthrough', level=1)
add_separator()
doc.add_paragraph()

# Step 1
doc.add_heading('Step 1: Extract Tableau Workbook', level=2)
add_body('The pipeline begins by extracting the .twb (XML) file from the .twbx package. A .twbx file is a ZIP archive containing the workbook XML and embedded data (Hyper files).')
add_bullet('Input: .twbx file (or .twb file directly)')
add_bullet('Output: Extracted .twb XML file in temp/ directory')
add_bullet('Module: parser/extractor.py')
add_body('If the input is already a .twb file, this step is skipped.')

doc.add_paragraph()

# Step 2
doc.add_heading('Step 2: Parse Tableau Metadata', level=2)
add_body('The XML parser reads the entire Tableau workbook structure and extracts 12 categories of metadata:')

add_table(
    ['Category', 'XML Source', 'What is extracted'],
    [
        ['Datasources', '<datasource> elements', 'Name, caption, connections, Hyper file paths'],
        ['Columns', '<column> + <metadata-record>', 'Name, caption, datatype, role, formula'],
        ['Calculations', '<column> with <calculation>', 'Calculated fields with formulas (non-trivial)'],
        ['Joins', '<relation type="join">', 'Left/right table, columns, join type'],
        ['Relationships', '<datasource-relationship>', 'Cross-datasource blending on shared columns'],
        ['Worksheets', '<worksheet>', 'Name, chart type, axes, filters, fields used'],
        ['Parameters', '<column param-domain-type>', 'Name, datatype, current value, allowable values'],
        ['Actions', '<action>', 'Name, type, source, target'],
        ['Dual Axis', '<worksheet> with multiple <axis>', 'Worksheets with dual-axis charts'],
        ['Table Calculations', 'WINDOW_ formulas', 'Tableau table calculations'],
        ['LOD Expressions', '{FIXED/INCLUDE/EXCLUDE}', 'Level-of-detail expressions'],
        ['Display Folders', '<folder> and <drill-path>', 'Field organization hierarchy'],
    ]
)

add_body('Key parsing decisions:')
add_bullet('Columns from both ', bold_prefix='<column> and <metadata-record>: ')
add_body('The <column> elements in Tableau XML contain calculated fields and dimension/measure definitions. However, many physical columns (from the data source) only appear in <metadata-record> elements. The parser reads both to ensure complete column coverage.')
add_bullet('Hidden columns are skipped (e.g., Order Quantity marked hidden="true")')
add_bullet('Tableau internal columns are skipped (e.g., :Measure Names, __tableau_internal)')
add_bullet('Number of Records (formula = "1") is skipped as it maps to COUNTROWS() in DAX')

doc.add_paragraph()

# Step 3
doc.add_heading('Step 3: Extract Data from Hyper Files', level=2)
add_body('Tableau stores data extracts in Hyper files (.hyper) inside the .twbx package. The tool extracts these and converts them to CSV format for Power BI consumption.')

add_body('Process:')
add_bullet('Extract .hyper files from the .twbx ZIP archive')
add_bullet('Map each Hyper file to its datasource caption using connection metadata')
add_bullet('Use the Tableau Hyper API (tableauhyperapi) to read tables')
add_bullet('Handle multiple schemas (public, Extract) - skip empty schemas')
add_bullet('Export each table as a CSV file named after the datasource caption')

add_body('Example mapping for the US Superstore workbook:')
add_table(
    ['Hyper File', 'Datasource Caption', 'CSV Output', 'Rows'],
    [
        ['Sales Planning newleaf.hyper', 'Sales Commission', 'Sales Commission.csv', '43'],
        ['dataengine_42019_618651678240lea.hyper', 'Sample - Superstore', 'Sample - Superstore.csv', '9,994'],
        ['dataengine_42019_622799629629lea.hyper', 'Sales Target', 'Sales Target.csv', '4,603'],
    ]
)

doc.add_paragraph()

# Step 4
doc.add_heading('Step 4: Generate Power BI Model (.bim)', level=2)
add_body('The BIM generator creates a complete Tabular Object Model (TOM) JSON file that defines the entire Power BI semantic model. This is the core output of the migration.')

add_body('The .bim file includes:')
add_bullet('Tables with M partitions', bold_prefix='Data Tables: ')
add_body('Each table has an M (Power Query) expression that loads data from the corresponding CSV file. The M expression includes column type definitions matching the original Tableau datatypes.')

add_bullet('Row-level formulas converted to DAX', bold_prefix='Calculated Columns: ')
add_body('Tableau calculated fields that operate at the row level (no aggregation functions) are converted to DAX calculated columns. These are evaluated for each row in the table.')

add_bullet('Aggregated formulas converted to DAX', bold_prefix='Measures: ')
add_body('Tableau calculated fields that use aggregation functions (SUM, AVG, COUNT, etc.) or LOD expressions are converted to DAX measures.')

add_bullet('Tableau parameters become a separate table with measures', bold_prefix='Parameters: ')
add_bullet('Datasource blending relationships with many-to-many cardinality', bold_prefix='Relationships: ')
add_bullet('Folder and drill-path hierarchy preserved as column properties', bold_prefix='Display Folders: ')

doc.add_paragraph()

# Step 5
doc.add_heading('Step 5: Generate Tabular Editor Scripts', level=2)
add_body('As a fallback deployment method, the tool generates C# scripts for Tabular Editor 2:')

add_table(
    ['Script', 'Purpose', 'When to Use'],
    [
        ['tabular_script.cs', 'Creates tables (calculated), calc columns, measures, parameters', 'When deploying via TE2 CLI to a blank PBI report'],
        ['measures_only.cs', 'Adds calc columns + measures to existing tables', 'When CSV data is loaded manually first'],
        ['relationships.cs', 'Creates many-to-many relationships', 'After saving the main script'],
        ['display_folders.cs', 'Sets DisplayFolder property on columns', 'After saving the main script'],
    ]
)

doc.add_paragraph()

# Step 6
doc.add_heading('Step 6: Deploy to Power BI Desktop', level=2)
add_body('The tool uses the PBIP (Power BI Project) format for deployment:')

add_bullet('Creates a PBIP folder structure with the .bim model', bold_prefix='Step 6a: ')
add_bullet('Opens the .pbir file in Power BI Desktop', bold_prefix='Step 6b: ')
add_bullet('PBI Desktop loads the model.bim directly with M partitions', bold_prefix='Step 6c: ')
add_bullet('User clicks Refresh to evaluate M expressions and load CSV data', bold_prefix='Step 6d: ')

add_body('PBIP Project Structure:')
add_code_block('''output/<workbook>/
  <workbook>.pbip                          # Project pointer file
  <workbook>.Report/
    definition.pbir                        # Report definition (points to semantic model)
    report.json                            # Minimal report layout
  <workbook>.SemanticModel/
    definition.pbism                       # Semantic model definition
    model.bim                              # Complete TOM model (tables, M partitions,
                                           #   columns, measures, relationships, folders)''')

doc.add_paragraph()

# Step 7
doc.add_heading('Step 7: Summary & Output', level=2)
add_body('The pipeline outputs a summary of all generated files and provides instructions for the remaining manual step (clicking Refresh in PBI Desktop to load data).')

doc.add_page_break()

# ===================================================================
# 4. MODULE REFERENCE
# ===================================================================
doc.add_heading('4. Module Reference', level=1)
add_separator()
doc.add_paragraph()

doc.add_heading('4.1 xml_parser.py - Functions', level=2)
add_table(
    ['Function', 'Description'],
    [
        ['load_xml(twb_path)', 'Parse TWB XML file and return root element'],
        ['get_datasources(root)', 'Extract datasource names, captions, connections, tables'],
        ['get_columns(root)', 'Extract columns from <column> and <metadata-record> elements'],
        ['get_calculations(root)', 'Extract calculated fields with formulas (skip trivial/parameters)'],
        ['get_joins(root)', 'Extract join definitions from <relation> elements'],
        ['get_relationships(root)', 'Extract datasource-level blending relationships'],
        ['get_worksheets(root)', 'Extract worksheet metadata (chart type, axes, filters)'],
        ['get_parameters(root)', 'Extract parameter definitions with current/allowable values'],
        ['get_display_folders(root)', 'Extract folder and drill-path hierarchy structures'],
        ['get_field_name_map(root)', 'Build internal name to caption mapping for all fields'],
        ['get_actions(root)', 'Extract dashboard actions'],
        ['get_dual_axis(root)', 'Detect worksheets with dual-axis configurations'],
        ['get_table_calculations(root)', 'Extract WINDOW_ table calculations'],
        ['get_lod_expressions(root)', 'Extract LOD expressions ({FIXED/INCLUDE/EXCLUDE})'],
    ]
)

doc.add_paragraph()

doc.add_heading('4.2 dax_converter.py - Conversion Rules', level=2)
add_table(
    ['Tableau Syntax', 'DAX Output', 'Example'],
    [
        ['IF cond THEN x ELSEIF cond2 THEN y ELSE z END', 'IF(cond, x, IF(cond2, y, z))', 'Nested IF() calls'],
        ['CASE [Field] WHEN val1 THEN r1 ... END', 'SWITCH([Field], val1, r1, ...)', 'SWITCH()'],
        ['{FIXED [Dim]: AGG([Meas])}', 'CALCULATE(AGG([Meas]), ALLEXCEPT(T, T[Dim]))', 'LOD expression'],
        ["DATEDIFF('day', [Start], [End])", 'DATEDIFF([Start], [End], DAY)', 'Date difference'],
        ['[Field]', "'Table'[Field]", 'Field reference with quoted table'],
        ['[DS].[Field]', "'Mapped DS'[Field]", 'Cross-datasource reference'],
        ['AVG([X])', 'AVERAGE([X])', 'Function mapping'],
        ['COUNTD([X])', 'DISTINCTCOUNT([X])', 'Function mapping'],
        ['NULL', 'BLANK()', 'Null handling'],
        ['INDEX()', '1 /* not supported */', 'Unsupported table calc'],
    ]
)

doc.add_page_break()

# ===================================================================
# 5. DAX CONVERSION ENGINE
# ===================================================================
doc.add_heading('5. DAX Conversion Engine', level=1)
add_separator()
doc.add_paragraph()

add_body('The DAX converter (dax_converter.py) processes Tableau formulas through a multi-stage pipeline:')

doc.add_heading('5.1 Conversion Pipeline', level=2)
add_bullet('Resolve internal field names to captions (Calculation_xxx to display name)', bold_prefix='Stage 1 - Name Resolution: ')
add_bullet('Convert {FIXED [Dim]: AGG([Meas])} to CALCULATE()', bold_prefix='Stage 2 - LOD Expressions: ')
add_bullet('Strip // comments from formulas', bold_prefix='Stage 3 - Comment Removal: ')
add_bullet('Convert [DS].[Field] to DS[Field] with datasource name mapping', bold_prefix='Stage 4 - Cross-Reference Resolution: ')
add_bullet("Convert standalone [Field] to 'Table'[Field] with proper quoting", bold_prefix='Stage 5 - Field Qualification: ')
add_bullet('Convert IF/ELSEIF/ELSE/END to nested IF() using recursive parser', bold_prefix='Stage 6 - IF Conversion: ')
add_bullet('Convert CASE/WHEN/THEN/END to SWITCH()', bold_prefix='Stage 7 - CASE Conversion: ')
add_bullet('Map AVG to AVERAGE, COUNTD to DISTINCTCOUNT, etc.', bold_prefix='Stage 8 - Function Mapping: ')
add_bullet('Replace NULL with BLANK()', bold_prefix='Stage 9 - NULL Handling: ')
add_bullet("Quote table names with special characters using single quotes ('Table Name')", bold_prefix='Stage 10 - Table Name Quoting: ')

doc.add_heading('5.2 Measure vs Calculated Column Classification', level=2)
add_body('The tool automatically classifies each Tableau calculated field as either a DAX measure or a DAX calculated column:')
add_bullet('Has aggregation functions (SUM, AVG, COUNT, MIN, MAX, COUNTD, ATTR)', bold_prefix='Measure: ')
add_bullet('Contains LOD expressions ({FIXED/INCLUDE})', bold_prefix='Measure: ')
add_bullet('No aggregation functions - operates at row level', bold_prefix='Calculated Column: ')

doc.add_page_break()

# ===================================================================
# 6. DATA FLOW DIAGRAM
# ===================================================================
doc.add_heading('6. Data Flow Diagram', level=1)
add_separator()
doc.add_paragraph()

add_code_block('''
  .twbx Input File
       |
       v
  +--------------------+
  | 1. Extract TWB     |  --> temp/<name>/<name>.twb
  +--------------------+
       |
       v
  +--------------------+
  | 2. Parse XML       |  --> Metadata Dictionary (13 categories)
  +--------------------+       |
       |                       +-> datasources, columns, calculations
       v                       +-> joins, relationships, worksheets
  +--------------------+       +-> parameters, actions, dual_axis
  | 3. Extract Hyper   |       +-> table_calcs, lods, display_folders
  |    Data            |       +-> field_name_map
  +--------------------+
       |
       v
  output/<name>/data/
    Sales Commission.csv (43 rows)
    Sample - Superstore.csv (9,994 rows)
    Sales Target.csv (4,603 rows)
       |
       v
  +--------------------+
  | 4. Generate .bim   |  --> output/<name>/model.bim
  +--------------------+       Contains:
       |                       - Tables with M partitions (CSV loading)
       v                       - Calculated columns (row-level DAX)
  +--------------------+       - Measures (aggregated DAX)
  | 5. Generate TE2    |       - Relationships (many-to-many)
  |    Scripts         |       - Display folders
  +--------------------+
       |
       v
  +--------------------+
  | 6. Create PBIP     |  --> output/<name>/<name>.pbip
  |    & Open PBI      |      output/<name>/<name>.Report/
  +--------------------+      output/<name>/<name>.SemanticModel/
       |
       v
  Power BI Desktop (with full model loaded)
''')

doc.add_page_break()

# ===================================================================
# 7. PREREQUISITES
# ===================================================================
doc.add_heading('7. Prerequisites & Environment Setup', level=1)
add_separator()
doc.add_paragraph()

doc.add_heading('7.1 System Requirements', level=2)
add_table(
    ['Requirement', 'Version', 'Purpose'],
    [
        ['Python', '3.8+', 'Runtime for the migration tool'],
        ['Power BI Desktop', '2.152+ (March 2026 or later)', 'Target platform - must support PBIP format'],
        ['Tabular Editor 2 (Portable)', '2.28.0+', 'Optional - for fallback script deployment'],
        ['Windows', '10/11', 'Required for Power BI Desktop'],
    ]
)

doc.add_heading('7.2 Python Libraries', level=2)
add_table(
    ['Library', 'Install Command', 'Purpose'],
    [
        ['tableauhyperapi', 'pip install tableauhyperapi', 'Read Tableau Hyper extract files'],
        ['python-docx', 'pip install python-docx', 'Generate this documentation (optional)'],
    ]
)

add_body('No other external Python libraries are required. The tool uses only Python standard library modules (xml.etree, json, zipfile, csv, os, subprocess, re, pathlib).')

doc.add_heading('7.3 Installation Steps', level=2)
add_body('1. Clone the repository:')
add_code_block('git clone https://github.com/anilthadathil/TabToPbi-Migrate.git\ncd TabToPbi-Migrate')

add_body('2. Install the required Python library:')
add_code_block('pip install tableauhyperapi')

add_body('3. (Optional) Place Tabular Editor 2 portable in a sibling directory:')
add_code_block('TabToPbi-Migrate/\n  migrate.py\n  parser/\n  samples/\nTabularEditor.2.28.0/\n  TabularEditor.exe')

add_body('4. Verify the installation:')
add_code_block('python migrate.py --help')

doc.add_page_break()

# ===================================================================
# 8. HOW TO RUN
# ===================================================================
doc.add_heading('8. How to Run the Tool', level=1)
add_separator()
doc.add_paragraph()

doc.add_heading('8.1 Basic Usage', level=2)
add_code_block('python migrate.py "path/to/workbook.twbx"')

add_body('Example:')
add_code_block('python migrate.py "samples/US_Superstore_10.0.twbx"')

doc.add_heading('8.2 What Happens Automatically', level=2)
add_bullet('Extracts TWB and Hyper data from the .twbx file')
add_bullet('Parses all Tableau metadata')
add_bullet('Converts Hyper data to CSV files')
add_bullet('Generates the .bim model with M expressions pointing to CSVs')
add_bullet('Creates a PBIP project folder')
add_bullet('Opens the project in Power BI Desktop')
add_bullet('Outputs a summary with file paths')

doc.add_heading('8.3 After the Script Completes', level=2)
add_bullet('Switch to Power BI Desktop')
add_bullet("Click Home > Refresh to load CSV data into the model")
add_bullet('Verify tables, columns, measures, and relationships in the Fields pane')
add_bullet("Click File > Save As to save as .pbix (click 'Don't upgrade' if TMDL prompt appears)")

doc.add_heading('8.4 Running from PowerShell', level=2)
add_code_block('cd "C:\\path\\to\\TabToPbi-Migrate"\npython migrate.py "C:\\path\\to\\workbook.twbx"')

doc.add_heading('8.5 Running on a Different Machine', level=2)
add_body('To run the tool on a new machine:')
add_bullet('Install Python 3.8+ from python.org')
add_bullet('Install Power BI Desktop (Microsoft Store or standalone)')
add_bullet('Clone the repository from GitHub')
add_bullet('Run: pip install tableauhyperapi')
add_bullet('(Optional) Download Tabular Editor 2 portable and place in sibling directory')
add_bullet('Run: python migrate.py "path/to/workbook.twbx"')

doc.add_page_break()

# ===================================================================
# 9. OUTPUT FILES
# ===================================================================
doc.add_heading('9. Output Files & Structure', level=1)
add_separator()
doc.add_paragraph()

add_code_block('''output/<workbook_name>/
  |
  +-- model.bim                    # Complete Power BI model (TOM JSON)
  +-- metadata.json                # Full Tableau metadata (for reference)
  |
  +-- data/
  |     Sales Commission.csv       # Extracted data tables
  |     Sample - Superstore.csv
  |     Sales Target.csv
  |
  +-- scripts/
  |     tabular_script.cs          # TE2 script: tables + columns + measures
  |     measures_only.cs           # TE2 script: measures only (for manual CSV load)
  |     relationships.cs           # TE2 script: relationships
  |     display_folders.cs         # TE2 script: folder organization
  |
  +-- <workbook_name>.pbip         # Power BI Project file
  +-- <workbook_name>.Report/
  |     definition.pbir            # Report definition
  |     report.json                # Report layout
  |
  +-- <workbook_name>.SemanticModel/
        definition.pbism           # Semantic model definition
        model.bim                  # Copy of the TOM model''')

doc.add_page_break()

# ===================================================================
# 10. LIMITATIONS
# ===================================================================
doc.add_heading('10. Limitations & Future Enhancements', level=1)
add_separator()
doc.add_paragraph()

doc.add_heading('10.1 Current Limitations', level=2)
add_table(
    ['Area', 'Limitation', 'Workaround'],
    [
        ['Data Refresh', 'User must click Refresh in PBI Desktop after opening', 'One-time manual step'],
        ['WINDOW_ Calcs', 'RUNNING_SUM, RANK, INDEX not converted', 'Manual DAX creation'],
        ['INCLUDE/EXCLUDE LOD', 'Only FIXED LOD is converted', 'Manual DAX creation'],
        ['Complex Date Functions', 'DATETRUNC, DATEADD, DATEPART not converted', 'Manual DAX creation'],
        ['Visual Migration', 'Worksheets/dashboards not migrated', 'Recreate visuals manually in PBI'],
        ['Live Connections', 'Only Hyper extracts are extracted', 'Connect to data source in PBI'],
        ['Nested Calculations', 'Calc fields referencing other calc fields may error', 'Check dependencies in PBI'],
        ['Save as .pbix', 'Must save manually via File > Save As', 'One-time manual step'],
    ]
)

doc.add_heading('10.2 Future Enhancements', level=2)
add_bullet('Automated data refresh via TMSL/XMLA after PBIP deployment')
add_bullet('Support for INCLUDE and EXCLUDE LOD expressions')
add_bullet('WINDOW_ table calculation conversion (RUNNING_SUM, RANK)')
add_bullet('Additional date function mappings (DATETRUNC, DATEADD, DATENAME)')
add_bullet('String function mappings (REPLACE, SPLIT, STARTSWITH)')
add_bullet('Visual migration (worksheet to PBI report page mapping)')
add_bullet('Batch processing of multiple workbooks')
add_bullet('Automated .pbix file creation and save')
add_bullet('Live connection support (SQL Server, PostgreSQL, etc.)')

doc.add_page_break()

# ===================================================================
# 11. TROUBLESHOOTING
# ===================================================================
doc.add_heading('11. Troubleshooting Guide', level=1)
add_separator()
doc.add_paragraph()

add_table(
    ['Issue', 'Cause', 'Solution'],
    [
        ['PBI Desktop not detected', 'No msmdsrv.exe process running', 'Open PBI Desktop with blank report first, then re-run'],
        ['Hyper data not extracted', 'tableauhyperapi not installed', 'Run: pip install tableauhyperapi'],
        ['DAX error: single value not determined', 'Column referenced without aggregation in measure', 'Check if calc should be a calculated column, not a measure'],
        ['DAX error: table name not resolved', "Table name has special characters", "Verify table names are quoted with single quotes in DAX"],
        ['TMDL upgrade prompt on save', 'PBI asks to convert to TMDL format', "Click 'Don't upgrade' to save in current format"],
        ['Fields pane empty after deploy', 'PBI UI not synced with backend', 'Close and reopen the PBIP project'],
        ['Circular dependency error', 'Calculated columns reference each other', 'Check formula dependencies and reorder'],
        ['Relationship error: duplicate values', 'One-to-many on non-unique column', 'Relationships use many-to-many cardinality by default'],
    ]
)

# ===================================================================
# SAVE
# ===================================================================
output_path = os.path.join('output', 'TabToPBI_Architecture_Document.docx')
os.makedirs('output', exist_ok=True)
doc.save(output_path)
print(f'Document saved: {os.path.abspath(output_path)}')
